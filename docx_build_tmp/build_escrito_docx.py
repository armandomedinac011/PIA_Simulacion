from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "ESCRITO_FINAL_PROYECTO_4_INTEGRANTES.md"
OUTPUT = ROOT / "deliverables" / "Escrito_Final_Proyecto_Simulacion_Equipo7.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(30, 38, 51)
MUTED = RGBColor(91, 103, 120)
LIGHT_FILL = "F2F4F7"
CALL_OUT_FILL = "F4F6F9"
BORDER = "DADCE0"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_border_bottom(paragraph, color="2E74B5", size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Title", 24, RGBColor(0, 0, 0), 0, 10),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    if "Code Block" not in [s.name for s in styles]:
        code = styles.add_style("Code Block", 1)
        code.font.name = "Courier New"
        code._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
        code._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
        code.font.size = Pt(9)
        code.font.color.rgb = RGBColor(40, 48, 58)
        code.paragraph_format.left_indent = Inches(0.18)
        code.paragraph_format.right_indent = Inches(0.05)
        code.paragraph_format.space_before = Pt(2)
        code.paragraph_format.space_after = Pt(2)
        code.paragraph_format.line_spacing = 1.0

    if "Callout" not in [s.name for s in styles]:
        callout = styles.add_style("Callout", 1)
        callout.font.name = "Calibri"
        callout._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        callout._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        callout.font.size = Pt(10.5)
        callout.font.color.rgb = RGBColor(31, 58, 95)
        callout.paragraph_format.left_indent = Inches(0.18)
        callout.paragraph_format.right_indent = Inches(0.18)
        callout.paragraph_format.space_before = Pt(6)
        callout.paragraph_format.space_after = Pt(8)
        callout.paragraph_format.line_spacing = 1.10


def set_headers_and_footers(doc):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(hp, after=0)
    run = hp.add_run("PIA Simulacion | Escrito final del proyecto")
    set_run_font(run, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(fp, after=0)
    run = fp.add_run("Documento para revision del equipo")
    set_run_font(run, size=9, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=4)
    run = p.add_run("PROYECTO INTEGRADOR ACADEMICO")
    set_run_font(run, size=11, color=MUTED, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(title, before=8, after=6, line=1.0)
    r = title.add_run("Simulacion estocastica de inventario y abastecimiento")
    set_run_font(r, size=24, color=RGBColor(0, 0, 0), bold=True)

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, after=16)
    r = subtitle.add_run("Centro de distribucion de materiales de construccion | Equipo 7")
    set_run_font(r, size=13, color=RGBColor(55, 55, 55))

    metadata = [
        ("Materia:", "Simulacion de Sistemas"),
        ("Documento:", "Escrito completo para compartir con integrantes del equipo"),
        ("Contenido:", "Contexto, codigo, Semana 3, Semana 4, Semana 5, Streamlit y Pygame"),
        ("Fecha:", datetime.now().strftime("%d/%m/%Y")),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=2)
        r1 = p.add_run(label + " ")
        set_run_font(r1, size=11, color=INK, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, size=11, color=INK)

    rule = doc.add_paragraph()
    add_border_bottom(rule)
    set_paragraph_spacing(rule, before=8, after=14)

    lead = doc.add_paragraph(style="Callout")
    add_inline(lead, "Este documento esta organizado para que cuatro personas puedan estudiar y exponer el proyecto. Incluye la explicacion del modelo, las variables, el codigo, la experimentacion, la optimizacion y las visualizaciones en Streamlit y Pygame.")
    shade_paragraph(lead, CALL_OUT_FILL)

    doc.add_paragraph("Contenido principal", style="Heading 1")
    contents = [
        "Descripcion general y contexto logistico.",
        "Variables, distribuciones y aleatoriedad de Semana 3.",
        "Motor de simulacion, eventos, KPIs y persistencia de datos.",
        "Experimentacion, optimizacion, reduccion de varianza y sensibilidad de Semana 4.",
        "Visualizacion final con Streamlit y Pygame para Semana 5.",
        "Analisis archivo por archivo y guia de exposicion para 4 integrantes.",
    ]
    for item in contents:
        p = doc.add_paragraph(style="List Bullet")
        add_inline(p, item)

    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=8, after=8)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_inline(paragraph, text):
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Courier New", size=9.5, color=DARK_BLUE)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def is_table_separator(line):
    stripped = line.strip()
    if not stripped.startswith("|"):
        return False
    cells = [cell.strip() for cell in stripped.strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        if not is_table_separator(lines[i]):
            rows.append([cell.strip() for cell in lines[i].strip().strip("|").split("|")])
        i += 1
    return rows, i


def column_widths(rows):
    cols = max(len(row) for row in rows)
    normalized = [row + [""] * (cols - len(row)) for row in rows]
    weights = []
    for idx in range(cols):
        values = [row[idx] for row in normalized]
        avg_len = sum(len(value) for value in values) / max(1, len(values))
        header_len = len(values[0]) if values else 1
        weights.append(max(1.2, min(4.0, avg_len / 18 + header_len / 22 + 0.8)))
    total = sum(weights)
    return [int(9360 * w / total) for w in weights]


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    normalized = [row + [""] * (cols - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = column_widths(normalized)
    set_table_width(table, widths)

    for r_idx, row in enumerate(normalized):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(paragraph, after=0, line=1.05)
            add_inline(paragraph, value)
            for run in paragraph.runs:
                set_run_font(run, size=9.2, bold=(r_idx == 0), color=INK)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=4)


def should_skip_source_heading(line, skipped_count):
    return skipped_count < 2 and line.startswith("# ")


def build_doc_from_markdown():
    doc = Document()
    style_document(doc)
    set_headers_and_footers(doc)
    add_cover(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    skipped_headings = 0
    code_lang = ""

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if should_skip_source_heading(line, skipped_headings):
            skipped_headings += 1
            i += 1
            continue

        if line.startswith("```"):
            in_code = not in_code
            code_lang = line.strip("`").strip()
            if in_code and code_lang:
                p = doc.add_paragraph(style="Code Block")
                add_inline(p, f"# {code_lang}")
                shade_paragraph(p, CALL_OUT_FILL)
            i += 1
            continue

        if in_code:
            p = doc.add_paragraph(style="Code Block")
            p.paragraph_format.keep_together = True
            text = line if line else " "
            run = p.add_run(text)
            set_run_font(run, name="Courier New", size=8.8, color=RGBColor(40, 48, 58))
            shade_paragraph(p, CALL_OUT_FILL)
            i += 1
            continue

        if not line.strip() or line.strip() == "---":
            i += 1
            continue

        if line.strip().startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            rows, next_i = parse_table(lines, i)
            add_table(doc, rows)
            i = next_i
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)), 3)
            text = heading.group(2).strip()
            style = "Heading 1" if level == 1 else "Heading 2" if level == 2 else "Heading 3"
            p = doc.add_paragraph(style=style)
            p.paragraph_format.keep_with_next = True
            add_inline(p, text)
            i += 1
            continue

        if line.startswith(">"):
            p = doc.add_paragraph(style="Callout")
            add_inline(p, line.lstrip("> ").strip())
            shade_paragraph(p, CALL_OUT_FILL)
            i += 1
            continue

        bullet = re.match(r"^\s*-\s+(.+)$", line)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            add_inline(p, bullet.group(1))
            i += 1
            continue

        number = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if number:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            add_inline(p, number.group(1))
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline(p, line.strip())
        i += 1

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc_from_markdown()
